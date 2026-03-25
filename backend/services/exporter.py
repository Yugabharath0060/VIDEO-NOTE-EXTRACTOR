"""
Exporter Service.
Exports extracted notes to PDF, Word (docx), and Markdown formats.
"""
import os
from datetime import datetime
from typing import Optional
from ..models.schemas import VideoNotes


def format_timestamp(seconds: float) -> str:
    """Convert seconds to HH:MM:SS format."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def export_markdown(notes: VideoNotes, output_path: str, include_timestamps: bool = True) -> str:
    """Export notes to a Markdown file."""
    lines = [
        f"# 🎬 {notes.title}",
        f"",
        f"**Source:** {notes.source}",
        f"**Extracted:** {datetime.now().strftime('%Y-%m-%d %H:%M')}",
    ]
    if notes.duration:
        lines.append(f"**Duration:** {format_timestamp(notes.duration)}")
    lines.append("")

    # Summary
    if notes.summary:
        lines += ["---", "", "## 📝 AI-Generated Notes", "", notes.summary, ""]

    # Transcript
    if notes.transcript:
        lines += ["---", "", "## 🗣️ Transcript", ""]
        for seg in notes.transcript:
            if include_timestamps:
                lines.append(f"**[{format_timestamp(seg.start)}]** {seg.text}")
            else:
                lines.append(seg.text)
        lines.append("")

    # OCR
    if notes.ocr_texts:
        lines += ["---", "", "## 🔍 Text Detected in Video (OCR)", ""]
        for ocr in notes.ocr_texts:
            if include_timestamps:
                lines.append(f"**[{format_timestamp(ocr.timestamp)}]** {ocr.text}")
            else:
                lines.append(f"- {ocr.text}")
        lines.append("")

    content = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    return output_path


def export_pdf(notes: VideoNotes, output_path: str, include_timestamps: bool = True) -> str:
    """Export notes to a PDF file using ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle", parent=styles["Title"],
        fontSize=22, spaceAfter=6,
        textColor=colors.HexColor("#1a1a2e")
    )
    heading_style = ParagraphStyle(
        "Heading2Custom", parent=styles["Heading2"],
        fontSize=14, spaceBefore=12, spaceAfter=4,
        textColor=colors.HexColor("#16213e")
    )
    body_style = ParagraphStyle(
        "BodyCustom", parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=4
    )
    ts_style = ParagraphStyle(
        "Timestamp", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#6c757d"), spaceAfter=2
    )

    # Title
    story.append(Paragraph(f"🎬 {notes.title}", title_style))
    story.append(Paragraph(f"Source: {notes.source}", ts_style))
    story.append(Paragraph(f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ts_style))
    if notes.duration:
        story.append(Paragraph(f"Duration: {format_timestamp(notes.duration)}", ts_style))
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dee2e6")))
    story.append(Spacer(1, 12))

    # Summary
    if notes.summary:
        story.append(Paragraph("📝 AI-Generated Notes", heading_style))
        for line in notes.summary.split("\n"):
            line = line.strip()
            if line:
                story.append(Paragraph(line.replace("#", "").replace("**", ""), body_style))
        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dee2e6")))
        story.append(Spacer(1, 12))

    # Transcript
    if notes.transcript:
        story.append(Paragraph("🗣️ Transcript", heading_style))
        for seg in notes.transcript:
            if include_timestamps:
                story.append(Paragraph(f"<font color='#6c757d'>[{format_timestamp(seg.start)}]</font> {seg.text}", body_style))
            else:
                story.append(Paragraph(seg.text, body_style))
        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dee2e6")))
        story.append(Spacer(1, 12))

    # OCR
    if notes.ocr_texts:
        story.append(Paragraph("🔍 Text Detected in Video (OCR)", heading_style))
        for ocr in notes.ocr_texts:
            if include_timestamps:
                story.append(Paragraph(f"<font color='#6c757d'>[{format_timestamp(ocr.timestamp)}]</font> {ocr.text}", body_style))
            else:
                story.append(Paragraph(f"• {ocr.text}", body_style))

    doc.build(story)
    return output_path


def export_word(notes: VideoNotes, output_path: str, include_timestamps: bool = True) -> str:
    """Export notes to a Word (.docx) file."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(f"🎬 {notes.title}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Source: {notes.source}\n").bold = True
    meta.add_run(f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if notes.duration:
        meta.add_run(f"\nDuration: {format_timestamp(notes.duration)}")

    doc.add_paragraph("─" * 60)

    # Summary
    if notes.summary:
        doc.add_heading("📝 AI-Generated Notes", level=1)
        for line in notes.summary.split("\n"):
            line = line.strip().replace("##", "").replace("**", "").replace("*", "")
            if line:
                doc.add_paragraph(line)
        doc.add_paragraph("─" * 60)

    # Transcript
    if notes.transcript:
        doc.add_heading("🗣️ Transcript", level=1)
        for seg in notes.transcript:
            p = doc.add_paragraph()
            if include_timestamps:
                ts_run = p.add_run(f"[{format_timestamp(seg.start)}] ")
                ts_run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
                ts_run.font.size = Pt(9)
            p.add_run(seg.text)
        doc.add_paragraph("─" * 60)

    # OCR
    if notes.ocr_texts:
        doc.add_heading("🔍 Text Detected in Video (OCR)", level=1)
        for ocr in notes.ocr_texts:
            p = doc.add_paragraph()
            if include_timestamps:
                ts_run = p.add_run(f"[{format_timestamp(ocr.timestamp)}] ")
                ts_run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
                ts_run.font.size = Pt(9)
            p.add_run(ocr.text)

    doc.save(output_path)
    return output_path


def export_notes(notes: VideoNotes, output_dir: str, fmt: str, include_timestamps: bool = True) -> str:
    """Main export dispatcher."""
    safe_title = "".join(c for c in notes.title if c.isalnum() or c in " _-")[:50].strip()
    filename = f"{notes.job_id}_{safe_title}"

    if fmt == "pdf":
        path = os.path.join(output_dir, f"{filename}.pdf")
        return export_pdf(notes, path, include_timestamps)
    elif fmt == "word":
        path = os.path.join(output_dir, f"{filename}.docx")
        return export_word(notes, path, include_timestamps)
    else:  # markdown
        path = os.path.join(output_dir, f"{filename}.md")
        return export_markdown(notes, path, include_timestamps)
